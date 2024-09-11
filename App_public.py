import telebot
import cohere
import docx
from gtts import gTTS
from apis import *
import fitz

# Configurações
bot = telebot.TeleBot(API_KEY_TELEGRAM)
co = cohere.Client(Cohere_API)

# Função para criar botões
def create_markup():
    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    help_button = telebot.types.KeyboardButton('HELP')
    markup.add(help_button)
    return markup

@bot.message_handler(commands=['start'])
def send_welcome(message):
    markup = create_markup()
    bot.send_message(message.chat.id, "Escolha uma opção:", reply_markup=markup)

@bot.message_handler(func=lambda message: message.text == 'HELP')
def send_help(message):
    help_text = (
        "Aqui estão os comandos disponíveis:\n"
        "/start - Inicia o bot\n"
        "/Help - Mostra esta mensagem de ajuda\n"
        "Resumir Documento - Envia um documento e recebe um resumo\n"
        "Enviar Áudio - Envia uma mensagem em áudio\n"
    )
    bot.send_message(message.chat.id, help_text)

# Função para gerar resumo de documentos usando Cohere
def resumir_doc(doc_path):
    doc = docx.Document(doc_path)
    texto_completo = []

    for c in doc.paragraphs:
        texto_completo.append(c.text)

    texto = '\n'.join(texto_completo)

    response = co.generate(
        model='command-xlarge-nightly',
        prompt=f"Faça um resumo detalhado e exclua o '*': {texto}",
        max_tokens=500
    )
    resumo = response.generations[0].text

    documento_resumido = docx.Document()
    documento_resumido.add_paragraph(resumo)

    resumo_path = "resumo_documento.docx"
    documento_resumido.save(resumo_path)

    return resumo_path

# Função para gerar resumo de PDFs
def resumir_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    texto_completo = []

    for pagina in doc:
        texto_completo.append(pagina.get_text())

    texto = '\n'.join(texto_completo)

    response = co.generate(
        model='command-xlarge-nightly',
        prompt=f"Faça um resumo com título e subtítulos: {texto}",
        max_tokens=3000
    )
    resumo = response.generations[0].text

    documento_resumido = docx.Document()
    documento_resumido.add_paragraph(resumo)

    resumo_path = "resumo_documento.docx"
    documento_resumido.save(resumo_path)

    return resumo_path

# Função para receber documentos
@bot.message_handler(content_types=['document'])
def handle_document(message):
    file_info = bot.get_file(message.document.file_id)
    downloaded_file = bot.download_file(file_info.file_path)

    with open(message.document.file_name, 'wb') as new_file:
        new_file.write(downloaded_file)

    if message.document.mime_type == 'application/pdf':
        resumo_path = resumir_pdf(message.document.file_name)
    else:
        resumo_path = resumir_doc(message.document.file_name)

    with open(resumo_path, 'rb') as doc:
        bot.send_document(message.chat.id, doc)
    bot.send_message(message.chat.id, "Resumo do documento enviado.")

# Função para áudio
@bot.message_handler(content_types=['voice'])
def open_audio(message):
    user_first_name = message.from_user.first_name
    user_last_name = message.from_user.last_name
    user_full_name = f"{user_first_name} {user_last_name}" if user_last_name else user_first_name

    texto = f"Oi {user_full_name}, no momento eu não consigo ouvir e responder a áudios, mas estou trabalhando nisso e em breve estarei pronto para isso!"
    tts = gTTS(text=texto, lang='pt', slow=False)

    arquivo_audio = "chatbot.mp3"
    tts.save(arquivo_audio)

    try:
        with open(arquivo_audio, 'rb') as audio_file:
            bot.send_audio(message.chat.id, audio_file)
    except Exception as e:
        bot.send_message(message.chat.id, "Falha ao enviar áudio.")

# Função para gerar respostas com Cohere AI
@bot.message_handler(func=lambda message: True)
def response(message):
    response = co.generate(
        model='command-xlarge-nightly',
        prompt=message.text,
        max_tokens=200
    )
    text_response = response.generations[0].text
    bot.reply_to(message, text_response)

# Inicia o bot
bot.polling()
