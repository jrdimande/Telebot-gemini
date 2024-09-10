import telebot
from apis import *
import google.generativeai as gemini
import webbrowser
import platform
import os
import random
import psutil
import pyautogui
import requests
import docx


# Config
bot = telebot.TeleBot(API_KEY_TELEGRAM)
gemini.configure(api_key=GEMINI_API_KEY)
model = gemini.GenerativeModel('gemini-pro')



# Função para criar botões
def create_markup():
    markup = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    chrome_button = telebot.types.KeyboardButton('Abrir Chrome')
    github_button = telebot.types.KeyboardButton('Abrir Github')
    brightness_button = telebot.types.KeyboardButton('Ajustar Brilho')
    system_button = telebot.types.KeyboardButton('Monitorar Sistema')
    vol_up_button = telebot.types.KeyboardButton('Volume +')
    vol_down_button = telebot.types.KeyboardButton('Volume -')
    screenshot_button = telebot.types.KeyboardButton('Captura')
    youtube_button = telebot.types.KeyboardButton("Abrir Youtube")
    markup.add(chrome_button, brightness_button, github_button, system_button, vol_down_button, vol_up_button,
               screenshot_button, youtube_button)

    return markup


@bot.message_handler(commands=['start'])
def send_welcome(message):
    markup = create_markup()
    bot.send_message(message.chat.id, "Escolha uma opção:", reply_markup=markup)


# Função para monitorar sistema
@bot.message_handler(func=lambda message: message.text == 'Monitorar Sistema')
def system_info(message):
    cpu_usage = psutil.cpu_percent(interval=1)
    memory_info = psutil.virtual_memory()
    storage = psutil.disk_usage('/')
    info = (f"Uso do CPU: {cpu_usage}%\nMemória Usada: {memory_info.used / (1024 ** 3): .2f} GB"
            f"\nMémoria Total: {memory_info.total / (1024 ** 3): .2f} GB\n"
            f"Armazenamento Usado: {storage.used / (1024 ** 3): .2f} GB\n"
            f"Armazenamento Total: {storage.total / (1024 ** 3): .2f} GB"
            f"\nArmazenamento Restante: {storage.total / (1024 ** 3) - storage.used / (1024 ** 3): .2f}")
    bot.reply_to(message, info)


# Função para aumentar volume
@bot.message_handler(func=lambda message: message.text == 'Volume +')
def volume_up(message):
    pyautogui.hotkey('volumeup')
    bot.send_message(message.chat.id, "Volume aumentado.")


# Função para diminuir volume
@bot.message_handler(func=lambda message: message.text == 'Volume -')
def volume_down(message):
    pyautogui.hotkey('volumedown')
    bot.send_message(message.chat.id, "Volume diminuído.")


# Função para fazer screenshot
@bot.message_handler(func=lambda message: message.text == 'Captura')
def take_screenshot(message):
    screenshot = pyautogui.screenshot()
    screenshot_path = "screenshot.png"
    screenshot.save(screenshot_path)

    with open(screenshot_path, "rb") as photo:
        bot.send_photo(message.chat.id, photo)
    bot.send_message(message.chat.id, "Captura de tela tirada e enviada.")


#####Funções para abrir aplicativos web#################################
@bot.message_handler(func=lambda message: message.text == 'Abrir Chrome')
def open_chrome(message):
    url = 'https://google.com'
    webbrowser.open(url)
    bot.send_message(message.chat.id, "Chrome aberto.")


@bot.message_handler(func=lambda message: message.text == 'Abrir Github')
def open_github(message):
    url = 'https://github.com'
    webbrowser.open(url)
    bot.send_message(message.chat.id, "Github aberto.")


@bot.message_handler(func=lambda message: message.text == 'Abrir Youtube')
def open_youtube(message):
    url = 'https://youtube.com'
    webbrowser.open(url)
    bot.reply_to(message, 'Youtube aberto')


# Função para ajustar brilho
@bot.message_handler(func=lambda message: message.text == 'Ajustar Brilho')
def adjust_brightness(message):
    levels = [20, 30, 40, 50, 60, 70, 80, 90, 100]
    level = random.choice(levels)

    try:
        if platform.system() == "Windows":
            os.system(
                f"powershell (Get-WmiObject -Namespace root/wmi -Class WmiMonitorBrightnessMethods).WmiSetBrightness(1,{level})")
            bot.send_message(message.chat.id, f"Brilho ajustado para {level}%.")
        else:
            bot.reply_to(message, "Ajuste de brilho não suportado neste sistema.")
    except Exception as e:
        bot.reply_to(message, f"{str(e)}")


# Função para gerar resumo de documentos
def resumir_doc(doc_path):
    doc  = docx.Document(doc_path)
    texto_completo = []

    for c in doc.paragraphs:
        texto_completo.append(c.text)

    texto = '\n'.join(texto_completo)

    response = model.generate_content(f"Faça resumo do texto com explicação: {texto}")
    resumo = response.candidates[0].content.parts[0].text

    documento_resumido = docx.Document()
    documento_resumido.add_paragraph(resumo)

    resumo_path = "resumo_documento.docx"
    documento_resumido.save(resumo_path)

    return resumo_path


# Função para receber documentos
@bot.message_handler(content_types=['document'])
def handle_document(message):
    file_info = bot.get_file(message.document.file_id)
    downloaded_file = bot.download_file(file_info.file_path)

    with open(message.document.file_name, 'wb') as new_file:
        new_file.write(downloaded_file)

    resumo_path = resumir_doc(message.document.file_name)

    with open(resumo_path, 'rb') as doc:
        bot.send_document(message.chat.id, doc)
    bot.send_message(message.chat.id, "Resumo do documento enviado.")


# Função para gerar respostas com AI
@bot.message_handler(func=lambda message: True)
def response(message):
    response = model.generate_content(message.text)
    text_response = response.candidates[0].content.parts[0].text
    bot.reply_to(message, text_response)


# Inicia o bot
bot.polling()
